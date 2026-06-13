import tkinter as tk
from tkinter import filedialog
import pandas as pd
import time
import numpy as np
from scipy.stats import skew
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import tempfile
import os

def add_dataframe_table(doc, df, title=None):
    """
    Add pandas DataFrame as formatted Word table.
    """

    if title:
        doc.add_heading(title, level=2)

    table = doc.add_table(
        rows=df.shape[0] + 1,
        cols=df.shape[1]
    )

    table.style = "Table Grid"

    # Header row
    for j, col in enumerate(df.columns):
        table.cell(0, j).text = str(col)

    # Data rows
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            table.cell(
                i + 1,
                j
            ).text = str(df.iloc[i, j])

    doc.add_paragraph()

def create_control_chart_fig(series, col):
    s = series.dropna()

    mean = s.mean()
    std = s.std()

    ucl = mean + 3 * std
    lcl = mean - 3 * std

    fig = plt.figure(figsize=(10, 4))
    plt.plot(s.index, s.values, label="Value")
    plt.axhline(mean, color="green", label="Mean")
    plt.axhline(ucl, color="red", linestyle="--", label="UCL")
    plt.axhline(lcl, color="red", linestyle="--", label="LCL")
    plt.title(f"Control Chart: {col}")
    plt.legend()

    plt.tight_layout()  

    return fig


# =========================================================
# TIMER START
# =========================================================
start_total = time.time()

# =========================================================
# ROOT WINDOW
# =========================================================
root = tk.Tk()
root.title("CSV EDA Tool")
root.geometry("450x600")

# =========================================================
# FILE PICKER
# =========================================================
file_path = filedialog.askopenfilename(
    title="Select Data File",
    filetypes=[
        ("Data files", "*.csv *.xlsx *.xls"),
        ("CSV files", "*.csv"),
        ("Excel files", "*.xlsx *.xls"),
        ("All files", "*.*")
    ]
)

if not file_path:
    raise ValueError("No file selected")

# =========================================================
# LOAD FILE
# =========================================================
print("\nLoading file...")

start_load = time.time()

file_ext = file_path.lower().split(".")[-1]

if file_ext == "csv":

    try:
        df = pd.read_csv(
            file_path,
            encoding="utf-8",
            sep=",",
            engine="python",
        )
    except UnicodeDecodeError:
        df = pd.read_csv(
            file_path,
            encoding="latin1",
            sep=",",
            engine="python",
        )

elif file_ext in ["xlsx", "xls"]:

    df = pd.read_excel(
        file_path
    )

else:
    raise ValueError(f"Unsupported file type: {file_ext}")

print(f"File loaded in {time.time() - start_load:.2f} seconds")

df.columns = df.columns.str.strip()

print("\nRows:", df.shape[0], "Columns:", df.shape[1])

# =========================================================
# STATE
# =========================================================
selected_columns = []
vars_dict = {}

# =========================================================
# FUNCTIONS
# =========================================================
def run_eda():
    global selected_columns

    selected_columns = [
        col for col, var in vars_dict.items()
        if var.get()
    ]

    if not selected_columns:
        print("No columns selected")
        return

    root.quit()


def select_all():
    for var in vars_dict.values():
        var.set(True)


def deselect_all():
    for var in vars_dict.values():
        var.set(False)

# =========================================================
# UI
# =========================================================
tk.Label(
    root,
    text="Select Columns for EDA",
    font=("Arial", 14, "bold")
).pack(pady=10)

container = tk.Frame(root)
container.pack(fill="both", expand=True)

canvas = tk.Canvas(container)
scrollbar = tk.Scrollbar(container, orient="vertical", command=canvas.yview)

scrollable_frame = tk.Frame(canvas)

scrollable_frame.bind(
    "<Configure>",
    lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
)

canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
canvas.configure(yscrollcommand=scrollbar.set)

canvas.pack(side="left", fill="both", expand=True)
scrollbar.pack(side="right", fill="y")

# =========================================================
# CHECKBOXES
# =========================================================
for col in df.columns:

    var = tk.BooleanVar(value=True)

    tk.Checkbutton(
        scrollable_frame,
        text=col,
        variable=var,
        anchor="w"
    ).pack(fill="x", padx=10, anchor="w")

    vars_dict[col] = var

# =========================================================
# BUTTONS
# =========================================================
button_frame = tk.Frame(root)
button_frame.pack(pady=10)

tk.Button(
    button_frame,
    text="Select All",
    command=select_all,
    width=15
).grid(row=0, column=0, padx=5)

tk.Button(
    button_frame,
    text="Deselect All",
    command=deselect_all,
    width=15
).grid(row=0, column=1, padx=5)

tk.Button(
    root,
    text="RUN EDA",
    command=run_eda,
    bg="green",
    fg="white",
    width=25,
    height=2
).pack(pady=10)

# =========================================================
# START UI
# =========================================================
root.mainloop()

# =========================================================
# VALIDATION
# =========================================================
if not selected_columns:
    raise ValueError("No columns selected")

eda_df = df[selected_columns]

# =========================================================
# BUILD ADVANCED EDA REPORT
# =========================================================
print("\nGenerating advanced EDA report...")

doc = Document()

doc.add_heading(
    "EDA REPORT FOR FEATURE ENGINEERING",
    level=1
)

doc.add_paragraph(
    f"Rows: {eda_df.shape[0]:,}"
)

doc.add_paragraph(
    f"Columns: {eda_df.shape[1]}"
)

# =========================================================
# DATA TYPES
# =========================================================
dtype_counts = (
    eda_df.dtypes
    .astype(str)
    .value_counts()
)

dtype_df = pd.DataFrame({
    "Data Type": dtype_counts.index,
    "Count": dtype_counts.values
})

add_dataframe_table(
    doc,
    dtype_df,
    "Data Types"
)

# =========================================================
# DATA QUALITY SCORECARD
# =========================================================

quality_rows = []

for col in eda_df.columns:

    s = eda_df[col]

    quality_rows.append({
        "Column": col,
        "Type": str(s.dtype),
        "Missing %": round(s.isnull().mean() * 100, 2),
        "Unique": s.nunique(),
        "Unique %": round(
            s.nunique() / len(s) * 100,
            2
        )
    })

quality_df = pd.DataFrame(quality_rows)

add_dataframe_table(
    doc,
    quality_df,
    "Data Quality Scorecard"
)

# =========================================================
# MISSING VALUES
# =========================================================


missing = (
    eda_df
    .isnull()
    .mean()
    .sort_values(ascending=False)
)

missing_df = pd.DataFrame({
    "Column": missing.index,
    "Missing %": (missing * 100).round(2).values
})

add_dataframe_table(
    doc,
    missing_df,
    "Missing Value Ranking"
)

# =========================================================
# CONSTANT FEATURES
# =========================================================
constant_cols = [
    c
    for c in eda_df.columns
    if eda_df[c].nunique(dropna=False) <= 1
]

constant_df = pd.DataFrame({
    "Column": constant_cols
})

if len(constant_df) == 0:
    constant_df = pd.DataFrame({
        "Column": ["None"]
    })

add_dataframe_table(
    doc,
    constant_df,
    "Constant Features"
)

# =========================================================
# NEAR CONSTANT FEATURES
# =========================================================
near_constant_rows = []

for col in eda_df.columns:

    vc = eda_df[col].astype(str).value_counts(
        normalize=True,
        dropna=False
    )

    if len(vc) > 0 and vc.iloc[0] > 0.95:

        near_constant_rows.append({
            "Column": col,
            "Dominant Value": str(vc.index[0]),
            "Frequency %": round(vc.iloc[0] * 100, 2)
        })

near_constant_df = pd.DataFrame(near_constant_rows)

if len(near_constant_df) == 0:
    near_constant_df = pd.DataFrame({
        "Column": ["None"]
    })

add_dataframe_table(
    doc,
    near_constant_df,
    "Near Constant Features"
)

# =========================================================
# NUMERIC ANALYSIS
# =========================================================
numeric_df = eda_df.select_dtypes(
    include=np.number
)
summary_rows = []

for col in numeric_df.columns:

    s = numeric_df[col]

    summary_rows.append({
        "Column": col,
        "Min": s.min(),
        "Max": s.max(),
        "Mean": round(s.mean(),4),
        "Median": round(s.median(),4),
        "Std": round(s.std(),4)
    })

summary_df = pd.DataFrame(summary_rows)

add_dataframe_table(
    doc,
    summary_df,
    "Numeric Feature Summary"
)

# =========================================================
# OUTLIERS
# =========================================================
outlier_rows = []

for col in numeric_df.columns:

    q1 = numeric_df[col].quantile(0.25)
    q3 = numeric_df[col].quantile(0.75)

    iqr = q3 - q1

    lower = q1 - 1.5 * iqr
    upper = q3 + 1.5 * iqr

    outliers = (
        (
            numeric_df[col] < lower
        ) |
        (
            numeric_df[col] > upper
        )
    ).sum()

    pct = (
        outliers /
        len(numeric_df)
    ) * 100

    outlier_rows.append({
        "Column": col,
        "Outliers": outliers,
        "Outlier %": round(pct, 2)
    })

outlier_df = pd.DataFrame(outlier_rows)

add_dataframe_table(
    doc,
    outlier_df,
    "Outlier Analysis"
)
# =========================================================
# SKEWNESS
# =========================================================
skew_rows = []

for col in numeric_df.columns:

    try:

        sk = skew(
            numeric_df[col].dropna()
        )

        if abs(sk) > 2:

            skew_rows.append({
                "Column": col,
                "Skewness": round(sk, 2)
            })

    except:
        pass

if len(skew_rows) == 0:

    skew_rows.append({
        "Column": "None",
        "Skewness": ""
    })

skew_df = pd.DataFrame(skew_rows)

add_dataframe_table(
    doc,
    skew_df,
    "Skewed Features"
)

# =========================================================
# CORRELATION ANALYSIS
# =========================================================
corr_rows = []

if numeric_df.shape[1] > 1:

    corr = (
        numeric_df
        .corr()
        .abs()
    )

    for i in range(len(corr.columns)):

        for j in range(i + 1, len(corr.columns)):

            val = corr.iloc[i, j]

            if val >= 0.95:

                corr_rows.append({
                    "Feature 1": corr.columns[i],
                    "Feature 2": corr.columns[j],
                    "Correlation": round(val, 3)
                })

if len(corr_rows) == 0:

    corr_rows.append({
        "Feature 1": "None",
        "Feature 2": "",
        "Correlation": ""
    })

corr_df = pd.DataFrame(corr_rows)

add_dataframe_table(
    doc,
    corr_df,
    "Highly Correlated Features"
)

# =========================================================
# CATEGORICAL ANALYSIS
# =========================================================
cat_cols = eda_df.select_dtypes(
    exclude=np.number
).columns

for col in cat_cols:

    top_vals = (
        eda_df[col]
        .astype(str)
        .value_counts()
        .head(10)
        .reset_index()
    )

    top_vals.columns = [
        "Value",
        "Count"
    ]

    add_dataframe_table(
        doc,
        top_vals,
        f"Categorical Feature: {col}"
    )

# =========================================================
# SAVE AS WORD REPORT (.DOCX)
# =========================================================

print("\nOpening save dialog...")

save_path = filedialog.asksaveasfilename(
    title="Save EDA Report",
    defaultextension=".docx",
    filetypes=[("Word Document", "*.docx")]
)

if save_path:

    print("Building Word report...")

    # -------------------------------------------------
    # TITLE
    # -------------------------------------------------

    doc.add_heading(
        "EDA REPORT FOR FEATURE ENGINEERING",
        level=1
    )

    # -------------------------------------------------
    # TEXT REPORT
    # -------------------------------------------------

    # -------------------------------------------------
    # CONTROL CHARTS
    # -------------------------------------------------

    if len(numeric_df.columns) > 0:

        doc.add_page_break()

        doc.add_heading(
            "Control Charts",
            level=1
        )

        for col in numeric_df.columns:

            s = numeric_df[col].dropna()

            if len(s) < 10:
                continue

            print(f"Adding chart: {col}")

            fig = create_control_chart_fig(s, col)

            temp_file = tempfile.NamedTemporaryFile(
                suffix=".png",
                delete=False
            )

            temp_name = temp_file.name

            temp_file.close()  

            fig.savefig(
                temp_name,
                dpi=300,
                bbox_inches="tight"
            )

            plt.close(fig)

            doc.add_heading(
                f"Control Chart: {col}",
                level=2
            )

            doc.add_picture(
                temp_name,
                width=Inches(6.5)
            )

    # -------------------------------------------------
    # SAVE
    # -------------------------------------------------

    doc.save(save_path)

    print(f"\nWord report saved to:\n{save_path}")

else:
    print("\nSave cancelled — report not saved")

print("\nEDA COMPLETE")